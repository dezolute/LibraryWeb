import os
from datetime import datetime
from typing import List

from docx import Document
from fastapi import HTTPException
from fastapi.responses import FileResponse
from starlette import status

from app.services.book import BookService
from app.services.reader import ReaderService
from app.repositories.sqlalchemy import SqlAlchemyRepository
from app.models import LoanORM
from app.models.types import BookAccessType, BookCopyStatus
from app.schemas import MultiDTO
from app.schemas.loan import LoanCreateDTO, LoanDTO
from app.schemas.relations import LoanRelationDTO
from app.schemas.utils import Pagination


class LoanService:
    def __init__(
            self,
            loan_repository: SqlAlchemyRepository[LoanORM],
            book_service: BookService,
            reader_service: ReaderService,
    ):
        self.loan_repository: SqlAlchemyRepository[LoanORM] = loan_repository
        self.book_service: BookService = book_service
        self.reader_service: ReaderService = reader_service

    async def get_loans(
        self,
        pg: Pagination,
        conditions = None,
        **filters
    ) -> MultiDTO[LoanRelationDTO]:
        loans, total = await self.loan_repository.find_all(
            pg=pg,
            conditions=conditions,
            **filters
        )

        loans_db = MultiDTO(
            items=[LoanRelationDTO.model_validate(loan) for loan in loans],
            total=total,
        )
        return loans_db

    async def create_loan(self, loan: LoanCreateDTO) -> LoanDTO:
        book = await self.book_service.get_single(id=loan.book_id)
        global copy
        for copy in book.copies:
            if copy.status == BookCopyStatus.RESERVED:
                await self.book_service.change_copy_status(
                    new_status=BookCopyStatus.BORROWED,
                    serial_num=copy.serial_num
                )
                break
        else:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Avaliable copy not found"
            )

        global db_loan

        if copy.access_type == BookAccessType.READING_ROOM:
            db_loan = await self.loan_repository.create({ 
                "reader_id": loan.reader_id,
                "copy_id": copy.serial_num,
                "due_date": datetime.now()
            })
        else:
            db_loan = await self.loan_repository.create({ 
                "reader_id": loan.reader_id,
                "copy_id": copy.serial_num
            })

        return LoanDTO.model_validate(db_loan)

    async def set_loan_as_returned(self, loan_id: int) -> LoanDTO:
        await self.loan_repository.update(data={ "return_date": datetime.now() }, id=loan_id)
        loan = await self.loan_repository.find(id=loan_id)


        await self.book_service.change_copy_status(
            new_status=BookCopyStatus.AVAILABLE,
            serial_num=loan.book_copy.serial_num
        )

        return LoanDTO.model_validate(loan)

    async def get_overdue_loans(self, pg: Pagination) -> MultiDTO[LoanRelationDTO]:
        loans, total = await self.loan_repository.find_all(
            pg=pg,
            conditions=[
                LoanORM.due_date < datetime.now(),
            ]
        )

        return MultiDTO(
            items=[
                LoanRelationDTO.model_validate(loan)
                for loan in loans
            ],
            total=total,
        )

    @staticmethod
    def create_report(loans: List[LoanRelationDTO], save_path: str) -> None:
        doc = Document()
        doc.add_heading("Отчёт по задолженности читателей", level=1)
        doc.add_paragraph(f"Дата создания: {datetime.today().strftime('%d.%m.%Y')}")

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ФИО"
        hdr_cells[1].text = "Книга"
        hdr_cells[2].text = "Серийный номер"
        hdr_cells[3].text = "Дата выдачи"
        hdr_cells[4].text = "Срок возврата"
        hdr_cells[5].text = "Просрочка (дн.)"

        for loan in loans:
            row_cells = table.add_row().cells
            row_cells[0].text = loan.reader.profile.full_name
            row_cells[1].text = loan.book_copy.book.title
            row_cells[2].text = loan.copy_id
            row_cells[3].text = loan.issue_date.strftime("%d.%m.%Y")
            row_cells[4].text = loan.due_date.strftime("%d.%m.%Y")
            row_cells[5].text = str((datetime.now() - loan.due_date).days)

        doc.save(save_path)

    async def overdue_report(self, pg: Pagination) -> FileResponse:
        loans = await self.get_overdue_loans(pg)

        path_to_file = os.path.join(os.path.abspath("."), "temp", f"overdue_report.docx")
        self.create_report(loans.items, path_to_file)

        return FileResponse(
            path=path_to_file,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{datetime.today().strftime('%d.%m.%Y')}_overdue.docx"
        )